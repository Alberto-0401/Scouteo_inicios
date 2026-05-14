# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Colores ────────────────────────────────────────────────────────────────
C_AZUL  = RGBColor(0x1E, 0x40, 0xAF)
C_GRIS  = RGBColor(0x6B, 0x72, 0x80)
C_NEGRO = RGBColor(0x11, 0x18, 0x27)
C_BLAN  = RGBColor(0xFF, 0xFF, 0xFF)
FONT    = "Calibri"

_bm_id = [0]

# ─── TOC ────────────────────────────────────────────────────────────────────
TOC = [
    (1,"1. ANALISIS DEL PROBLEMA","s1"),
    (2,"   1.1. Introduccion","s1_1"),
    (2,"   1.2. Objetivos principales","s1_2"),
    (2,"   1.3. Funciones y rendimientos deseados","s1_3"),
    (2,"   1.4. Planteamiento y evaluacion de soluciones","s1_4"),
    (2,"   1.5. Justificacion de la solucion elegida","s1_5"),
    (2,"   1.6. Modelado de la solucion","s1_6"),
    (3,"      1.6.1. Recursos humanos","s1_6_1"),
    (3,"      1.6.2. Recursos hardware","s1_6_2"),
    (3,"      1.6.3. Recursos software","s1_6_3"),
    (2,"   1.7. Planificacion temporal","s1_7"),
    (1,"2. DISENO E IMPLEMENTACION DEL PROYECTO","s2"),
    (2,"   2.1. Plan de empresa","s2_1"),
    (2,"   2.2. Arquitectura de la aplicacion","s2_2"),
    (2,"   2.3. Tipos de usuarios y roles","s2_3"),
    (2,"   2.4. Mapa de navegacion","s2_4"),
    (2,"   2.5. Diseno de la base de datos","s2_5"),
    (3,"      2.5.1. Diagrama Entidad-Relacion","s2_5_1"),
    (3,"      2.5.2. Modelo relacional (tablas)","s2_5_2"),
    (2,"   2.6. Desarrollo por modulos","s2_6"),
    (3,"      2.6.1. Acceso a Datos","s2_6_1"),
    (3,"      2.6.2. Sistemas de Gestion Empresarial","s2_6_2"),
    (3,"      2.6.3. Desarrollo de Interfaces","s2_6_3"),
    (3,"      2.6.4. Programacion de Servicios y Procesos","s2_6_4"),
    (3,"      2.6.5. Programacion Multimedia y Dispositivos Moviles","s2_6_5"),
    (1,"3. FASE DE PRUEBAS","s3"),
    (1,"4. DOCUMENTACION DE LA APLICACION","s4"),
    (2,"   4.1. Introduccion a la aplicacion","s4_1"),
    (2,"   4.2. Manual de instalacion","s4_2"),
    (2,"   4.3. Manual de usuario","s4_3"),
    (2,"   4.4. Manual de administracion","s4_4"),
    (1,"5. CONCLUSIONES FINALES","s5"),
    (1,"6. BIBLIOGRAFIA","s6"),
]

# ─── Helpers ────────────────────────────────────────────────────────────────
def add_bookmark(para, name):
    bid = str(_bm_id[0]); _bm_id[0] += 1
    p = para._p
    bs = OxmlElement("w:bookmarkStart"); bs.set(qn("w:id"),bid); bs.set(qn("w:name"),name)
    be = OxmlElement("w:bookmarkEnd"); be.set(qn("w:id"),bid)
    p.insert(0, bs); p.append(be)

def toc_entry(doc, text, level, bookmark=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm((level - 1) * 0.7)
    # right-aligned tab with dot leader at right margin
    pPr = p._p.get_or_add_pPr()
    tabs_el = OxmlElement('w:tabs')
    ts = OxmlElement('w:tab')
    ts.set(qn('w:val'), 'right'); ts.set(qn('w:leader'), 'dot'); ts.set(qn('w:pos'), '8900')
    tabs_el.append(ts); pPr.append(tabs_el)
    font_size = Pt(11 if level == 1 else 10.5); bold = (level == 1)
    run = p.add_run(text.strip())
    run.font.name = FONT; run.font.size = font_size; run.font.bold = bold; run.font.color.rgb = C_NEGRO
    if bookmark:
        r_tab = p.add_run(); tab_el = OxmlElement('w:tab'); r_tab._r.append(tab_el)
        _add_pageref(p, bookmark, font_size, bold)

def page_break(doc):
    p = doc.add_paragraph(); run = p.add_run()
    br = OxmlElement("w:br"); br.set(qn("w:type"),"page"); run._r.append(br)

def set_bottom_border(para, color="1E40AF", sz="8"):
    pPr = para._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),sz)
    bot.set(qn("w:space"),"4"); bot.set(qn("w:color"),color)
    pBdr.append(bot); pPr.append(pBdr)

def set_para_shade(para, fill="F3F4F6"):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),fill)
    pPr.append(shd)

def set_cell_bg(cell, hex6):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex6)
    tcPr.append(shd)

def set_line_spacing(para, spacing=1.15):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = spacing

def body(doc, text, bold=False, italic=False, center=False, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    set_line_spacing(p)
    run = p.add_run(text)
    run.font.name = FONT; run.font.size = Pt(12)
    run.font.bold = bold; run.font.italic = italic
    run.font.color.rgb = color if color else C_NEGRO
    return p

def note(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    set_para_shade(p, "F0F4FF")
    set_line_spacing(p)
    run = p.add_run(text)
    run.font.name = FONT; run.font.size = Pt(11)
    run.font.italic = True; run.font.color.rgb = RGBColor(0x1E,0x40,0xAF)

def bul(doc, text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5 + indent*0.4)
    set_line_spacing(p)
    run = p.add_run(text)
    run.font.name = FONT; run.font.size = Pt(12); run.font.color.rgb = C_NEGRO

def code_block(doc, text):
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        set_para_shade(p, "F3F4F6")
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"; run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x1F,0x29,0x37)

def make_table(doc, headers, rows, col_widths=None, stripe=True):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c = hr.cells[i]; set_cell_bg(c,"1E40AF")
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold=True; r.font.color.rgb=C_BLAN
        r.font.name=FONT; r.font.size=Pt(10)
    for ri,row in enumerate(rows):
        dr = t.rows[ri+1]
        bg = ("EBF2FF" if ri%2==0 else "FFFFFF") if stripe else "FFFFFF"
        for ci,val in enumerate(row):
            c = dr.cells[ci]; set_cell_bg(c,bg)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name=FONT; r.font.size=Pt(10); r.font.color.rgb=C_NEGRO
    if col_widths:
        for i,w in enumerate(col_widths):
            for row in t.rows: row.cells[i].width=Cm(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)
    return t

def figure_caption(doc, num, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    run = p.add_run(f"Figura {num}: {text}")
    run.font.name=FONT; run.font.size=Pt(10); run.font.italic=True; run.font.color.rgb=C_GRIS

def insert_figure(doc, img_path, num, caption_text, width_cm=15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    figure_caption(doc, num, caption_text)

def h1(doc, text, bookmark):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24); p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_bottom_border(p)
    run = p.add_run(text.upper())
    run.font.name=FONT; run.font.size=Pt(14); run.font.bold=True; run.font.color.rgb=C_AZUL
    add_bookmark(p, bookmark); return p

def h2(doc, text, bookmark):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name=FONT; run.font.size=Pt(12); run.font.bold=True; run.font.color.rgb=C_NEGRO
    add_bookmark(p, bookmark); return p

def h3(doc, text, bookmark):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name=FONT; run.font.size=Pt(12); run.font.bold=True; run.font.italic=True
    run.font.color.rgb=RGBColor(0x1E,0x40,0xAF)
    add_bookmark(p, bookmark); return p

def _add_pageref(p, bookmark, font_size, bold):
    r1 = p.add_run()
    fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin')
    r1._r.append(fc)
    r2 = p.add_run()
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = f' PAGEREF {bookmark} '; r2._r.append(instr)
    r3 = p.add_run()
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'separate')
    r3._r.append(fc2)
    r4 = p.add_run('?')
    r4.font.name = FONT; r4.font.size = font_size; r4.font.bold = bold; r4.font.color.rgb = C_NEGRO
    r5 = p.add_run()
    fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'end')
    r5._r.append(fc3)

def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.font.name=FONT; run.font.size=Pt(10); run.font.color.rgb=C_GRIS
        for tag, val in [("begin",None),("instrText","PAGE"),("end",None)]:
            fc = OxmlElement("w:fldChar" if tag in ("begin","end") else "w:instrText")
            if tag in ("begin","end"):
                fc.set(qn("w:fldCharType"),tag)
            else:
                fc.text = val
            run._r.append(fc)

# ════════════════════════════════════════════════════════════════════════
doc = Document()
for sec in doc.sections:
    sec.top_margin=Cm(2.54); sec.bottom_margin=Cm(2.54)
    sec.left_margin=Cm(2.54); sec.right_margin=Cm(2.54)
    sec.page_width=Cm(21); sec.page_height=Cm(29.7)

add_page_numbers(doc)

# ── PORTADA ───────────────────────────────────────────────────────────────
for _ in range(7): doc.add_paragraph()

logo = doc.add_paragraph()
logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = logo.add_run("SCOUTEO")
r.font.name=FONT; r.font.size=Pt(52); r.font.bold=True; r.font.color.rgb=C_AZUL

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Sistema de Gestion Deportiva para Clubes de Futbol")
r2.font.name=FONT; r2.font.size=Pt(16); r2.font.color.rgb=C_GRIS

doc.add_paragraph()
sep = doc.add_paragraph(); sep.alignment=WD_ALIGN_PARAGRAPH.CENTER
set_bottom_border(sep,"1E40AF","12")
doc.add_paragraph()

for txt,sz,bold,color in [
    ("Proyecto Intermodular",14,True,C_NEGRO),
    ("CFGS Desarrollo de Aplicaciones Multiplataforma (DAM)",13,False,C_GRIS),
    ("",10,False,C_GRIS),
    ("Alumno: Alberto Perez Oncina",12,True,C_NEGRO),
    ("Tutor: [Nombre del tutor]",12,False,C_GRIS),
    ("Centro: [Nombre del centro educativo]",12,False,C_GRIS),
    ("Curso academico: 2024 - 2025",12,False,C_GRIS),
]:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(txt); r.font.name=FONT; r.font.size=Pt(sz); r.font.bold=bold; r.font.color.rgb=color

page_break(doc)

# ── INDICE ────────────────────────────────────────────────────────────────
idx_t = doc.add_paragraph()
idx_t.paragraph_format.space_before=Pt(0); idx_t.paragraph_format.space_after=Pt(12)
r=idx_t.add_run("TABLA DE CONTENIDO")
r.font.name=FONT; r.font.size=Pt(16); r.font.bold=True; r.font.color.rgb=C_AZUL
set_bottom_border(idx_t)

for level, text, anchor in TOC:
    toc_entry(doc, text, level, anchor)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# 1. ANALISIS DEL PROBLEMA
# ════════════════════════════════════════════════════════════════════════
h1(doc,"1. ANALISIS DEL PROBLEMA","s1")

# 1.1 Introduccion
h2(doc,"1.1. Introduccion","s1_1")
body(doc,
    "Scouteo es un sistema de gestion deportiva destinado a clubes de futbol de categorias "
    "inferiores. El proyecto surge de la necesidad identificada en clubes de base, donde la "
    "gestion de jugadores, partidos, entrenamientos y datos medicos se realiza habitualmente "
    "con herramientas no especializadas: hojas de calculo, grupos de mensajeria instantanea "
    "y cuadernos en papel.")
body(doc,
    "El sistema se compone de dos elementos: un cliente de escritorio multiplataforma "
    "desarrollado con JavaFX y una API REST desplegada en Google Cloud Run que actua como "
    "backend centralizado. Los datos se almacenan en una base de datos PostgreSQL gestionada "
    "en Supabase, lo que garantiza la disponibilidad de la informacion desde cualquier "
    "instalacion del cliente.")
body(doc,
    "El proyecto abarca los modulos del ciclo formativo de grado superior en Desarrollo de "
    "Aplicaciones Multiplataforma (DAM): Acceso a Datos, Desarrollo de Interfaces, "
    "Programacion de Servicios y Procesos, Programacion Multimedia y Dispositivos Moviles, "
    "y Sistemas de Gestion Empresarial.")

# 1.2 Objetivos
h2(doc,"1.2. Objetivos principales","s1_2")
body(doc,"Los objetivos principales del proyecto son los siguientes:")
for obj in [
    "Desarrollar una aplicacion de escritorio multiplataforma en JavaFX que permita la gestion integral de clubes de futbol.",
    "Implementar una API REST con Spring Boot desplegada en la nube que actue como capa de acceso a datos para todos los clientes.",
    "Centralizar la persistencia de datos en una base de datos PostgreSQL remota (Supabase), garantizando la coherencia y disponibilidad.",
    "Proporcionar herramientas de analisis estadistico y generacion de informes en PDF mediante JasperReports.",
    "Garantizar la seguridad del sistema mediante autenticacion JWT y cifrado BCrypt para las contrasenas.",
    "Implementar un sistema de roles (directiva, entrenador, jugador) que controle el acceso a las funcionalidades.",
    "Distribuir la aplicacion cliente como instalador nativo para Windows mediante jpackage.",
]:
    bul(doc, obj)

# 1.3 Funciones y rendimientos
h2(doc,"1.3. Funciones y rendimientos deseados","s1_3")
body(doc,"A continuacion se describen los servicios que ofrece el sistema una vez implementado:")

make_table(doc,
    ["Modulo funcional","Descripcion","Rol/es que lo usan"],
    [
        ["Gestion de jugadores","Alta, edicion, baja, cambio de estado (activo/lesionado/cedido) y foto de perfil","Directiva, Entrenador"],
        ["Gestion de equipos","Creacion de equipos por categoria y temporada, asignacion de entrenadores","Directiva"],
        ["Registro de partidos","Resultado, estadisticas colectivas (posesion, tiros, corners) y formacion tactica","Directiva, Entrenador"],
        ["Convocatorias","Seleccion de titulares y suplentes por partido","Entrenador"],
        ["Estadisticas individuales","Goles, asistencias, tarjetas, minutos y valoracion por jugador y partido","Directiva, Entrenador"],
        ["Entrenamientos","Programacion de sesiones y registro de asistencias","Entrenador"],
        ["Historial medico","Registro de lesiones con fechas de recuperacion estimada y real","Directiva, Entrenador"],
        ["Objetivos","Metas individuales y colectivas con porcentaje de progreso","Directiva, Entrenador"],
        ["Ranking","Clasificacion dinamica de goleadores y asistentes con filtros","Todos"],
        ["Informes PDF","Generacion de reportes de plantilla, partido y temporada","Directiva, Entrenador"],
        ["Graficos de rendimiento","KPIs y graficos de evolucion con TilesFX","Directiva, Entrenador"],
        ["Gestion de club","Configuracion de datos del club y temporada activa","Directiva"],
        ["Gestion de usuarios","Alta de entrenadores y jugadores como usuarios del sistema","Directiva"],
    ],
    col_widths=[4.5, 7.5, 3.5])

# 1.4 Soluciones
h2(doc,"1.4. Planteamiento y evaluacion de diversas soluciones","s1_4")
body(doc,
    "Se han evaluado tres alternativas tecnologicas antes de seleccionar la arquitectura "
    "definitiva del proyecto.")

body(doc,"Solucion A: Aplicacion web (SPA + Node.js)", bold=True)
body(doc,
    "Consistiria en un frontend de pagina unica desarrollado con React o Vue.js y un "
    "backend en Node.js con Express. La base de datos seria accesible desde cualquier "
    "navegador sin necesidad de instalacion.")
make_table(doc,["Ventajas","Inconvenientes"],[
    ["Accesible desde cualquier dispositivo con navegador","Curva de aprendizaje elevada para React/Vue (no cubierta en el ciclo)"],
    ["Sin instalacion en el cliente","El ciclo DAM no cubre desarrollo web frontend de forma profunda"],
    ["Actualizaciones inmediatas sin redistribuir","Menor integracion con los modulos del ciclo"],
],col_widths=[7.5,8])

body(doc,"Solucion B: Aplicacion de escritorio JavaFX + API REST (SELECCIONADA)", bold=True)
body(doc,
    "Un cliente de escritorio en JavaFX 25 comunica con una API REST desarrollada en "
    "Spring Boot. La base de datos PostgreSQL se aloja en Supabase. Esta solucion integra "
    "directamente los modulos de Acceso a Datos, Desarrollo de Interfaces y Programacion "
    "de Servicios y Procesos del ciclo.")
make_table(doc,["Ventajas","Inconvenientes"],[
    ["Tecnologia Java estudiada en profundidad en el ciclo DAM","Requiere instalacion en cada equipo cliente"],
    ["Cubre todos los modulos del ciclo (DI, AD, PSP, PMM)","Actualizaciones requieren redistribuir el instalador"],
    ["Interfaz rica y nativa con JavaFX (FXML, CSS)",""],
    ["API REST reutilizable por cualquier otro cliente (movil, web)",""],
    ["Generacion de PDF profesional con JasperReports",""],
],col_widths=[7.5,8])

body(doc,"Solucion C: Aplicacion movil nativa Android (Kotlin)", bold=True)
body(doc,
    "Una aplicacion Android desarrollada en Kotlin con Retrofit para las llamadas HTTP y "
    "Room como base de datos local. Se orientaria exclusivamente a dispositivos moviles.")
make_table(doc,["Ventajas","Inconvenientes"],[
    ["Accesible desde smartphones en cualquier lugar","Pantalla reducida, inadecuada para gestion compleja"],
    ["Notificaciones push nativas","Solo compatible con Android (no multiplataforma)"],
    ["Experiencia de usuario mobile-first","La gestion de clubes requiere vistas de escritorio"],
],col_widths=[7.5,8])

# 1.5 Justificacion
h2(doc,"1.5. Justificacion de la solucion elegida","s1_5")
body(doc,
    "Se ha seleccionado la Solucion B por los siguientes motivos:")
for r in [
    "Java y JavaFX son las tecnologias centrales del modulo de Desarrollo de Interfaces del ciclo, por lo que su uso maximiza la coherencia con el proyecto formativo.",
    "La separacion cliente-servidor mediante API REST permite que la logica de negocio quede centralizada en el backend, facilitando el mantenimiento y la escalabilidad futura.",
    "Supabase proporciona una base de datos PostgreSQL gestionada en la nube de forma gratuita para proyectos de desarrollo, eliminando la necesidad de infraestructura propia.",
    "Spring Boot cubre el modulo de Acceso a Datos al proporcionar un ORM (Spring Data JPA / Hibernate) y un sistema de migraciones compatible con el esquema relacional disenado.",
    "El uso de HttpClient nativo de Java y la programacion multihilo mediante JavaFX Task cubre el modulo de Programacion de Servicios y Procesos.",
    "La generacion de un instalador nativo con jpackage cubre el criterio de distribucion del modulo de Desarrollo de Interfaces.",
    "La API REST es consumida tambien por la aplicacion movil complementaria (Scouteo Mobile), cubriendo asi el modulo de Programacion Multimedia y Dispositivos Moviles.",
]:
    bul(doc, r)

# 1.6 Modelado
h2(doc,"1.6. Modelado de la solucion","s1_6")

h3(doc,"1.6.1. Recursos humanos","s1_6_1")
make_table(doc,["Rol","Persona","Responsabilidades"],[
    ["Desarrollador principal","Alberto Perez Oncina","Analisis, diseno, desarrollo, pruebas y documentacion del proyecto completo"],
    ["Tutor academico","[Nombre del tutor]","Supervision, orientacion y evaluacion del proyecto"],
],col_widths=[4,4,7.5])

h3(doc,"1.6.2. Recursos hardware","s1_6_2")
make_table(doc,["Recurso","Especificaciones","Uso"],[
    ["PC de desarrollo","Intel Core i7, 16 GB RAM, SSD 512 GB, Windows 11 Pro","Desarrollo, pruebas locales y generacion del instalador"],
    ["Servidor backend","Google Cloud Run (instancias serverless 2 vCPU / 2 GB RAM)","Ejecucion de la API REST en produccion"],
    ["Base de datos","Supabase Free Tier (PostgreSQL 15, 500 MB almacenamiento)","Persistencia de datos en la nube"],
    ["Repositorio de codigo","GitHub (repositorio privado)","Control de versiones y despliegue continuo"],
],col_widths=[3.5,5.5,6.5])

h3(doc,"1.6.3. Recursos software","s1_6_3")
make_table(doc,["Herramienta","Version","Proposito"],[
    ["JDK (Amazon Corretto)","24","Compilacion y ejecucion del cliente JavaFX y del backend Spring Boot"],
    ["JavaFX","25","Framework de interfaz grafica para la aplicacion de escritorio"],
    ["Spring Boot","3.3","Framework para el desarrollo de la API REST del backend"],
    ["Gradle","8.12","Sistema de construccion del proyecto cliente JavaFX"],
    ["Maven","3.9","Sistema de construccion del proyecto backend Spring Boot"],
    ["Supabase","Cloud","Base de datos PostgreSQL gestionada en la nube"],
    ["IntelliJ IDEA","2024.1","IDE principal para Java / Spring Boot"],
    ["Visual Studio Code","1.89","IDE para configuracion, scripts y documentacion"],
    ["Git / GitHub","2.47","Control de versiones y repositorio remoto"],
    ["Docker","26","Contenerizacion de la API para el despliegue en Cloud Run"],
    ["Postman","11","Pruebas manuales de los endpoints de la API REST"],
    ["DBeaver","24","Gestion y consulta de la base de datos Supabase"],
    ["JasperReports","7.0.3","Generacion de informes PDF desde el cliente JavaFX"],
    ["ControlsFX","11.2.1","Componentes UI avanzados para JavaFX"],
    ["TilesFX","21.0.9","Graficos e indicadores tipo dashboard para JavaFX"],
    ["BCrypt (favre-lib)","0.10.2","Cifrado de contrasenas en el backend"],
    ["Gson","2.10.1","Serializacion/deserializacion JSON en el cliente"],
],col_widths=[4,2.5,9])

# 1.7 Planificacion
h2(doc,"1.7. Planificacion temporal","s1_7")
body(doc,
    "La planificacion del proyecto se ha distribuido a lo largo del segundo y tercer "
    "trimestre del curso academico 2024-2025. La Figura 1 representa el diagrama de Gantt "
    "aproximado del proyecto.")

note(doc,"Figura 1: Diagrama de Gantt del proyecto — se recomienda insertar aqui el diagrama generado con Notion, Trello o herramienta equivalente.")

make_table(doc,
    ["Fase","Tareas","Oct","Nov","Dic","Ene","Feb","Mar","Abr","May"],
    [
        ["Analisis","Requisitos, casos de uso, diseno E/R","X","X","","","","","",""],
        ["Diseno BD","Esquema Supabase, migraciones SQL","","X","X","","","","",""],
        ["Backend API","Endpoints REST, Spring Security, JWT","","","X","X","","","",""],
        ["Cliente - Base","Login, Dashboard, CRUD jugadores/equipos","","","","X","X","","",""],
        ["Cliente - Avanzado","Partidos, entrenamientos, estadisticas","","","","","X","X","",""],
        ["Informes y graficos","JasperReports, TilesFX, ranking","","","","","","X","X",""],
        ["Pruebas","Bateria de pruebas, correcciones","","","","","","","X","X"],
        ["Instalador y docs","jpackage, documentacion tecnica","","","","","","","","X"],
    ],
    col_widths=[2.5,5.5,1,1,1,1,1,1,1,1])

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# 2. DISENO E IMPLEMENTACION
# ════════════════════════════════════════════════════════════════════════
h1(doc,"2. DISENO E IMPLEMENTACION DEL PROYECTO","s2")

# 2.1 Plan de empresa
h2(doc,"2.1. Plan de empresa","s2_1")
body(doc,
    "A efectos del proyecto, se plantea la empresa ficticia Scouteo Solutions S.L., una "
    "startup tecnologica especializada en software de gestion para el deporte base.")

make_table(doc,["Campo","Descripcion"],[
    ["Razon social","Scouteo Solutions S.L."],
    ["Actividad","Desarrollo y comercializacion de software de gestion deportiva"],
    ["Sector","Sports tech / Software as a Service (SaaS)"],
    ["Mercado objetivo","Clubes de futbol de categorias infantil, cadete, juvenil y aficionado en Espana"],
    ["Modelo de negocio","Licencia anual por club (50-150 EUR/ano segun numero de equipos) + soporte tecnico"],
    ["Propuesta de valor","Centralizacion digital de la gestion del club, eliminando el uso de WhatsApp y Excel como herramientas de gestion"],
    ["Competidores","Sportymanager, Dugout, Footix (soluciones web); Scouteo se diferencia por su cliente de escritorio de alto rendimiento y su integracion de informes PDF"],
    ["Estrategia de distribucion","Instalador gratuito descargable; API en la nube con subscription mensual"],
],col_widths=[4.5,11])

body(doc,
    "El proyecto se enmarca en los modulos de Sistemas de Gestion Empresarial al plantear "
    "un sistema que cubre funciones similares a un CRM deportivo: gestion de clientes "
    "(jugadores y familias), seguimiento de actividad (partidos y entrenamientos), y "
    "generacion de informes para la toma de decisiones por parte de la directiva.")

# 2.2 Arquitectura
h2(doc,"2.2. Arquitectura de la aplicacion","s2_2")
body(doc,
    "El sistema sigue una arquitectura cliente-servidor de tres capas. El cliente JavaFX "
    "aplica el patron MVC (Modelo-Vista-Controlador): las vistas se definen en archivos "
    "FXML, los controladores Java gestionan la logica de la interfaz, y los modelos son "
    "POJOs que representan las entidades del dominio. La comunicacion con el servidor se "
    "delega en la capa DAO a traves del componente ApiClient.")
body(doc,
    "El backend implementa una arquitectura en capas propia de Spring Boot: controladores "
    "REST, servicios con logica de negocio, repositorios Spring Data JPA y entidades "
    "mapeadas a las tablas de Supabase. La seguridad se gestiona con Spring Security y "
    "un filtro JWT que valida el token en cada peticion.")

insert_figure(doc, r"C:\Users\User\Downloads\arquitectura.drawio.png", 2, "Diagrama de arquitectura del sistema")

code_block(doc,"""
+----------------------------------------------+
|          CLIENTE (Escritorio JavaFX)         |
|                                              |
|  Vista FXML <-> Controlador <-> DAO          |
|                       |                      |
|               ApiClient (Singleton)          |
|           HttpClient nativo Java 11+         |
+---------------------|------------------------+
                       | HTTPS + JWT
                       v
+----------------------------------------------+
|       BACKEND  (Spring Boot / Cloud Run)     |
|                                              |
|  @RestController                             |
|       |                                      |
|  @Service  (logica de negocio)               |
|       |                                      |
|  @Repository  (Spring Data JPA)              |
|       |                                      |
|  Supabase  (PostgreSQL 15)                   |
+----------------------------------------------+
                       ^
                       | HTTPS + JWT
+----------------------------------------------+
|     CLIENTE MOVIL  (Scouteo Mobile)          |
|     React Native / Android                   |
+----------------------------------------------+""")

# 2.3 Roles
h2(doc,"2.3. Tipos de usuarios y roles","s2_3")
body(doc,
    "El sistema define tres roles de usuario con distintos niveles de acceso a las "
    "funcionalidades de la aplicacion:")
make_table(doc,
    ["Funcionalidad","Directiva","Entrenador","Jugador"],
    [
        ["Ver todos los equipos del club","Si","No (solo los suyos)","No"],
        ["Crear y eliminar equipos","Si","No","No"],
        ["Gestionar jugadores (CRUD)","Si","Si (su equipo)","No"],
        ["Registrar partidos y estadisticas","Si","Si (su equipo)","No"],
        ["Gestionar convocatorias","Si","Si (su equipo)","No"],
        ["Gestionar entrenamientos","Si","Si (su equipo)","No"],
        ["Gestionar historial medico","Si","Si (su equipo)","No"],
        ["Ver estadisticas y ranking","Si","Si (su equipo)","Si (las propias)"],
        ["Generar informes PDF","Si","Si","No"],
        ["Configurar datos del club","Si","No","No"],
        ["Gestionar usuarios del sistema","Si","No","No"],
        ["Ver su propia ficha","Si","Si","Si"],
    ],
    col_widths=[7,2.2,2.8,2.5])
body(doc,
    "La logica de control de acceso se implementa en dos niveles: en el backend mediante "
    "anotaciones de Spring Security (@PreAuthorize), y en el cliente JavaFX ocultando o "
    "desactivando los elementos de la interfaz segun el rol almacenado en SesionUsuario.")

# 2.4 Mapa de navegacion
h2(doc,"2.4. Mapa de navegacion","s2_4")
body(doc,
    "La estructura de navegacion de la aplicacion de escritorio parte de una pantalla de "
    "autenticacion y conduce a un panel principal (Dashboard) que actua como hub central. "
    "Desde el Dashboard se accede a todos los modulos a traves de un menu lateral.")

insert_figure(doc, r"C:\Users\User\Downloads\navegacion.drawio.png", 3, "Mapa de navegacion de la aplicacion")

make_table(doc,
    ["Pantalla","Pantallas accesibles desde ella","Rol minimo requerido"],
    [
        ["Login.fxml","Dashboard (si autenticacion correcta)","Todos"],
        ["Dashboard.fxml","Jugadores, Equipos, Partidos, Entrenamientos, Historial, Objetivos, Ranking, Estadisticas, Graficos, Informes, InfoClub, Configuracion","Todos (segun rol)"],
        ["ListadoJugadores.fxml","FormJugador (alta/edicion), EstadisticasJugador","Directiva, Entrenador"],
        ["FormJugador.fxml","(Vuelve a ListadoJugadores)","Directiva, Entrenador"],
        ["Partidos.fxml","FormPartido (alta/edicion), Convocatorias, FormEstadisticasPartido","Directiva, Entrenador"],
        ["FormPartido.fxml","(Vuelve a Partidos)","Directiva, Entrenador"],
        ["Entrenamientos.fxml","(Registro de asistencias inline)","Entrenador, Directiva"],
        ["EstadisticasJugador.fxml","GraficosRendimiento","Directiva, Entrenador"],
        ["Informes.fxml","(Generacion PDF a disco)","Directiva, Entrenador"],
        ["Loading.fxml","Login (splash inicial)","Sistema"],
    ],
    col_widths=[4,9,2.5])

# 2.5 Base de datos
h2(doc,"2.5. Diseno de la base de datos","s2_5")

h3(doc,"2.5.1. Diagrama Entidad-Relacion","s2_5_1")
body(doc,
    "El modelo de datos se compone de 13 entidades principales. A continuacion se "
    "describen las relaciones entre ellas:")

make_table(doc,
    ["Entidad","Relacion","Entidad relacionada","Cardinalidad"],
    [
        ["Club","tiene","Equipo","1:N"],
        ["Club","tiene","Usuario","1:N"],
        ["Club","tiene","Configuracion","1:1"],
        ["Equipo","pertenece a","Club","N:1"],
        ["Equipo","tiene","Jugador","1:N"],
        ["Equipo","juega","Partido","1:N"],
        ["Equipo","realiza","Entrenamiento","1:N"],
        ["Equipo","asignado a (via EquipoEntrenador)","Usuario (rol entrenador)","N:M"],
        ["Jugador","participa en (via Alineacion)","Partido","N:M"],
        ["Jugador","asiste a (via AsistenciaEntrenamiento)","Entrenamiento","N:M"],
        ["Jugador","tiene","HistorialMedico","1:N"],
        ["Jugador","tiene","Objetivo","1:N"],
        ["Partido","genera","EventoPartido","1:N"],
        ["Jugador","tiene en partido (via JugadorPartido)","EstadisticaPartido","1:N"],
    ],
    col_widths=[3.5,4,4,2.5])

insert_figure(doc, r"C:\Users\User\Downloads\E_R.drawio.png", 4, "Diagrama Entidad-Relacion", width_cm=16)

h3(doc,"2.5.2. Modelo relacional (tablas)","s2_5_2")
body(doc,"Las principales tablas del modelo relacional en Supabase son las siguientes:")

for tabla, pk, fks, campos in [
    ("clubs","id (PK)","--","nombre, creado_en"),
    ("usuarios","id (PK)","club_id (FK clubs)","email, password_hash, rol, nombre, apellidos, telefono, foto_url, activo, ultimo_acceso"),
    ("equipos","id (PK)","club_id (FK clubs)","nombre, categoria, temporada, campo, escudo_url, activo"),
    ("equipo_entrenador","id (PK)","equipo_id (FK equipos), usuario_id (FK usuarios)","--"),
    ("jugadores","id (PK)","equipo_id (FK equipos), usuario_id (FK usuarios, nullable)","nombre, apellidos, fecha_nacimiento, dorsal, posicion, pierna_dominante, altura_cm, peso_kg, foto_url, estado"),
    ("partidos","id (PK)","equipo_id (FK equipos)","fecha_hora, rival, tipo, campo, superficie, competicion, temporada, formacion, goles_favor, goles_contra, posesion_pct, tiros_totales, tiros_a_puerta, corners, faltas, notas"),
    ("alineaciones","id (PK)","partido_id (FK partidos), jugador_id (FK jugadores)","minuto_entrada, minuto_salida, posicion_partido, goles, asistencias, tarjetas_amarillas, tarjetas_rojas, paradas, goles_encajados, porteria_cero, valoracion, nota_entrenador"),
    ("entrenamientos","id (PK)","equipo_id (FK equipos)","fecha_hora, duracion_min, tipo, intensidad, objetivos, observaciones"),
    ("asistencias_entrenamiento","id (PK)","entrenamiento_id (FK entrenamientos), jugador_id (FK jugadores)","estado, motivo"),
    ("historial_medico","id (PK)","jugador_id (FK jugadores)","tipo_lesion, zona_afectada, fecha_lesion, fecha_recuperacion_est, fecha_alta, observaciones"),
    ("objetivos","id (PK)","equipo_id (FK equipos), jugador_id (FK jugadores, nullable)","descripcion, fecha_inicio, fecha_limite, prioridad, estado, progreso_pct, observaciones"),
    ("eventos_partido","id (PK)","partido_id (FK partidos), jugador_id (FK jugadores)","minuto, tipo, jugador2_id, descripcion"),
    ("configuracion","id (PK)","club_id (FK clubs)","nombre_club, localidad, presidente, email, telefono, escudo, temporada_actual"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    run = p.add_run(tabla)
    run.font.name="Courier New"; run.font.size=Pt(11); run.font.bold=True; run.font.color.rgb=C_AZUL
    code_block(doc, f"  PK: {pk}\n  FK: {fks}\n  Campos: {campos}")
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

# 2.6 Modulos
h2(doc,"2.6. Desarrollo por modulos","s2_6")
body(doc,
    "A continuacion se describe como el proyecto aborda cada uno de los modulos "
    "intermodulares del ciclo formativo DAM.")

h3(doc,"2.6.1. Acceso a Datos","s2_6_1")
body(doc,
    "Este modulo se cubre en los dos extremos de la arquitectura: el backend y el cliente.")
body(doc,"En el backend (Scouteo-API):", bold=True)
for r in [
    "Se utiliza Spring Data JPA con Hibernate como ORM para el mapeo de entidades Java a tablas PostgreSQL en Supabase.",
    "Los repositorios extienden JpaRepository, que proporciona los metodos CRUD basicos y permite definir consultas personalizadas mediante @Query.",
    "Se aplica el patron DTO (Data Transfer Object) para desacoplar la representacion interna de las entidades de la API publica.",
    "La conexion a Supabase se configura en application-prod.properties mediante variables de entorno inyectadas en Cloud Run.",
]:
    bul(doc, r)

body(doc,"En el cliente (Scouteo-JavaFX):", bold=True)
for r in [
    "Se implementa el patron DAO (Data Access Object) con 13 clases que encapsulan las llamadas a la API REST.",
    "La comunicacion HTTP se realiza mediante HttpClient nativo de Java 11+, sin dependencias externas.",
    "La serializacion y deserializacion de JSON se realiza con Gson, incluyendo adaptadores personalizados para LocalDate y LocalDateTime.",
    "El componente ApiClient es un Singleton que gestiona el token JWT y lo aniade automaticamente a cada peticion en el header Authorization.",
]:
    bul(doc, r)

body(doc,"Clase ApiClient — fragmento de codigo:", bold=False, italic=True)
code_block(doc,"""private String request(String method, String path, String bodyJson) {
    HttpRequest.Builder builder = HttpRequest.newBuilder()
        .uri(URI.create(BASE_URL + path))
        .header("Content-Type", "application/json");
    if (jwtToken != null)
        builder.header("Authorization", "Bearer " + jwtToken);
    // ... construccion y envio de la peticion
    HttpResponse<String> response = httpClient.send(builder.build(),
        HttpResponse.BodyHandlers.ofString(StandardCharsets.UTF_8));
    if (response.statusCode() >= 200 && response.statusCode() < 300)
        return response.body();
    return null;
}""")

h3(doc,"2.6.2. Sistemas de Gestion Empresarial","s2_6_2")
body(doc,
    "Scouteo actua como un sistema de gestion empresarial (CRM) especializado en el "
    "ambito deportivo. Los elementos que cubren este modulo son:")
for r in [
    "Gestion de la entidad Club como unidad organizativa central con su configuracion propia (nombre, presidente, contacto, temporada activa).",
    "Sistema de roles y permisos que define que operaciones puede realizar cada tipo de usuario, equivalente al control de acceso de un ERP.",
    "Gestion de recursos humanos del club: jugadores como activos del equipo con su estado, historial medico y ficha tecnica.",
    "Generacion de informes para la toma de decisiones: informe de plantilla, informe de rendimiento por partido e informe de temporada.",
    "Seguimiento de objetivos con indicadores de progreso (KPIs), similar a los cuadros de mando de un sistema ERP.",
]:
    bul(doc, r)

h3(doc,"2.6.3. Desarrollo de Interfaces","s2_6_3")
body(doc,
    "La interfaz de usuario de la aplicacion de escritorio se ha desarrollado con JavaFX 25 "
    "siguiendo los principios de usabilidad y accesibilidad:")
for r in [
    "Las vistas se definen declarativamente en 20 archivos FXML, separando completamente el diseno de la logica.",
    "Los estilos visuales se gestionan mediante CSS (scouteo.css), lo que permite modificar la apariencia sin tocar el codigo Java.",
    "Se utilizan componentes avanzados de ControlsFX para mejorar la experiencia de usuario (autocomplete, validacion de formularios, notificaciones).",
    "Los graficos de rendimiento se implementan con TilesFX, que proporciona tiles animados tipo dashboard.",
    "Los informes PDF se generan con JasperReports 7.0.3, a partir de plantillas .jrxml que separan el diseno del informe de los datos.",
    "La aplicacion se distribuye como instalador nativo para Windows mediante jpackage, incluyendo el JRE necesario para su ejecucion.",
    "Se implementa una pantalla de carga animada (LoadingController + FootballLoadingPane) que mejora la percepcion de rendimiento.",
]:
    bul(doc, r)

h3(doc,"2.6.4. Programacion de Servicios y Procesos","s2_6_4")
body(doc,
    "El modulo de Programacion de Servicios y Procesos se cubre mediante:")

body(doc,"Programacion multihilo en el cliente JavaFX:", bold=True)
body(doc,
    "Las operaciones de red (llamadas HTTP a la API) se ejecutan en threads separados del "
    "hilo principal de JavaFX (FX Application Thread) para evitar el bloqueo de la "
    "interfaz. Se utilizan dos mecanismos segun la complejidad de la operacion:")
for r in [
    "Thread anonimo + Platform.runLater(): para operaciones simples que necesitan actualizar la UI al completarse.",
    "JavaFX Task<T>: para operaciones con progreso reportable, vinculadas a indicadores de carga visibles en la pantalla.",
]:
    bul(doc, r)

body(doc,"Fragmento de codigo — carga asincrona de jugadores:", bold=False, italic=True)
code_block(doc,"""new Thread(() -> {
    List<Jugador> jugadores = jugadorDAO.obtenerPorEquipo(equipoId);
    Platform.runLater(() -> {
        tableJugadores.getItems().setAll(jugadores);
        loadingPane.setVisible(false);
    });
}).start();""")

body(doc,"Comunicacion en red (HTTP sobre HTTPS):", bold=True)
for r in [
    "Toda la comunicacion cliente-servidor se realiza mediante peticiones HTTP/HTTPS utilizando HttpClient de Java 11+.",
    "El cliente establece una conexion persistente (keep-alive) con el servidor para reducir la latencia en operaciones frecuentes.",
    "Se implementa manejo de errores de red: timeouts, reintentos y notificacion al usuario cuando la API no responde.",
    "La autenticacion sin estado (stateless) mediante JWT elimina la necesidad de mantener sesiones en el servidor.",
]:
    bul(doc, r)

h3(doc,"2.6.5. Programacion Multimedia y Dispositivos Moviles","s2_6_5")
body(doc,
    "El modulo de Programacion Multimedia y Dispositivos Moviles se cubre mediante "
    "la aplicacion complementaria Scouteo Mobile, disponible en el repositorio "
    "Scouteo-Mobile, y mediante los componentes multimedia integrados en el cliente "
    "de escritorio.")

body(doc,"Scouteo Mobile:", bold=True)
for r in [
    "Aplicacion movil complementaria que consume la misma API REST que el cliente de escritorio.",
    "Orientada a jugadores y entrenadores para consultar convocatorias, estadisticas personales y proximos partidos desde el movil.",
    "Implementa una interfaz adaptada a pantallas pequenas con navegacion por gestos.",
    "Se comunica con la API REST mediante peticiones HTTP autenticadas con JWT, reutilizando la misma capa de seguridad del backend.",
]:
    bul(doc, r)

body(doc,"Componentes multimedia en el cliente de escritorio:", bold=True)
for r in [
    "TilesFX proporciona animaciones fluidas en los tiles del dashboard de estadisticas.",
    "Los jugadores y equipos pueden tener fotos de perfil/escudos almacenadas en Supabase Storage y visualizadas en la interfaz.",
    "La pantalla de carga (FootballLoadingPane) incluye una animacion CSS personalizada.",
]:
    bul(doc, r)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# 3. PRUEBAS
# ════════════════════════════════════════════════════════════════════════
h1(doc,"3. FASE DE PRUEBAS","s3")
body(doc,
    "Se ha elaborado una bateria de pruebas que cubre los principales flujos de la "
    "aplicacion, tanto para detectar errores logicos como para verificar el rendimiento "
    "en condiciones normales de uso.")

make_table(doc,
    ["ID","Modulo","Descripcion de la prueba","Resultado esperado","Resultado obtenido","Estado"],
    [
        ["P-01","Autenticacion","Inicio de sesion con credenciales correctas","Navegacion al Dashboard y almacenamiento del JWT en SesionUsuario","Correcto","PASS"],
        ["P-02","Autenticacion","Inicio de sesion con contrasena incorrecta","Mensaje de error visible; no se navega al Dashboard","Correcto","PASS"],
        ["P-03","Autenticacion","Acceso a endpoint protegido sin token","API devuelve 401 Unauthorized","Correcto","PASS"],
        ["P-04","Jugadores","Alta de jugador con dorsal duplicado en el equipo","Mensaje de validacion que impide el guardado","Correcto","PASS"],
        ["P-05","Jugadores","Alta de jugador con todos los datos validos","Jugador visible en la tabla con datos correctos","Correcto","PASS"],
        ["P-06","Jugadores","Edicion del estado de un jugador (activo -> lesionado)","El estado cambia en la tabla y en Supabase","Correcto","PASS"],
        ["P-07","Partidos","Registro de partido con resultado y estadisticas colectivas","Partido visible en la lista con datos correctos","Correcto","PASS"],
        ["P-08","Partidos","Registro de estadisticas individuales por jugador","Estadisticas reflejadas en el ranking y perfil del jugador","Correcto","PASS"],
        ["P-09","Convocatorias","Seleccion de titular y suplentes para un partido","Alineacion guardada correctamente en la API","Correcto","PASS"],
        ["P-10","Entrenamientos","Registro de asistencia: asistio, falta justificada, falta injustificada","Estado de cada jugador guardado correctamente","Correcto","PASS"],
        ["P-11","Historial medico","Alta de lesion con fecha de recuperacion estimada","Lesion visible en el historial; jugador marcado como lesionado","Correcto","PASS"],
        ["P-12","Ranking","Cambio de criterio de ordenacion (goles -> asistencias)","La tabla se actualiza correctamente con el nuevo criterio","Correcto","PASS"],
        ["P-13","Informes","Generacion de informe de plantilla en PDF","PDF generado correctamente con los datos del equipo","Correcto","PASS"],
        ["P-14","Roles","Acceso de entrenador al modulo de equipos (no permitido)","Opcion oculta/desactivada en el menu lateral","Correcto","PASS"],
        ["P-15","Red","Uso de la aplicacion sin conexion a Internet","Mensaje de aviso de API no disponible; modo offline parcial","Correcto","PASS"],
        ["P-16","Rendimiento","Carga de lista de 100 jugadores","Tiempo de carga inferior a 2 segundos","1.3 segundos","PASS"],
        ["P-17","Multihilo","Operacion de red sin bloqueo de la UI","La interfaz permanece responsiva durante la carga de datos","Correcto","PASS"],
        ["P-18","Seguridad","Peticion a la API con token expirado","API devuelve 401; la aplicacion redirige al login","Correcto","PASS"],
        ["P-19","Instalador","Instalacion en un equipo limpio (sin JDK instalado)","La aplicacion se instala y ejecuta correctamente","Correcto","PASS"],
        ["P-20","PDF","Informe de partido con alineacion completa","PDF con datos de todos los jugadores convocados","Correcto","PASS"],
    ],
    col_widths=[1,2.5,4,3.5,2.5,1.5])

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# 4. DOCUMENTACION
# ════════════════════════════════════════════════════════════════════════
h1(doc,"4. DOCUMENTACION DE LA APLICACION","s4")

h2(doc,"4.1. Introduccion a la aplicacion","s4_1")
body(doc,
    "Scouteo es una aplicacion de escritorio para Windows que permite a los clubes de "
    "futbol gestionar de forma centralizada toda la informacion relacionada con sus "
    "equipos, jugadores y actividad deportiva. La aplicacion requiere conexion a Internet "
    "para sincronizar los datos con el servidor en la nube.")
body(doc,
    "La interfaz se organiza en torno a un panel principal con menu lateral desde el que "
    "se accede a todos los modulos. El selector de equipo en la parte superior permite "
    "cambiar el contexto entre los distintos equipos del club sin cerrar la sesion.")

h2(doc,"4.2. Manual de instalacion","s4_2")
body(doc,"Requisitos del sistema:", bold=True)
make_table(doc,["Requisito","Valor minimo"],[
    ["Sistema operativo","Windows 10 64-bit (o superior)"],
    ["Procesador","Intel Core i5 de 8a generacion o equivalente"],
    ["Memoria RAM","4 GB (recomendado 8 GB)"],
    ["Espacio en disco","500 MB libres"],
    ["Conexion a Internet","Requerida (mínimo 5 Mbps)"],
    ["JDK","No requerido (incluido en el instalador mediante jpackage)"],
],col_widths=[5,10.5])

body(doc,"Proceso de instalacion:", bold=True)
for n,paso in enumerate([
    "Descargar el archivo ScouteoSetup.exe desde el repositorio o enlace proporcionado por el centro.",
    "Ejecutar ScouteoSetup.exe con permisos de administrador.",
    "Seguir el asistente de instalacion (aceptar terminos, elegir directorio de destino).",
    "Al finalizar, hacer clic en 'Iniciar Scouteo' o acceder desde el acceso directo del escritorio.",
    "En el primer inicio, introducir las credenciales de directiva proporcionadas por el administrador del club.",
],1):
    bul(doc, f"Paso {n}: {paso}")

h2(doc,"4.3. Manual de usuario","s4_3")
body(doc,"A continuacion se describe el uso de las principales funcionalidades de la aplicacion.")

body(doc,"Inicio de sesion:", bold=True)
body(doc,"Se introduce el correo electronico y la contrasena del usuario en la pantalla de login. Si las credenciales son correctas, la aplicacion navega al Dashboard y muestra el panel principal del club.")

body(doc,"Gestion de jugadores:", bold=True)
for paso in [
    "Acceder al modulo 'Jugadores' desde el menu lateral.",
    "Seleccionar el equipo en el desplegable superior.",
    "Hacer clic en 'Nuevo jugador' para dar de alta un jugador, o seleccionar uno existente para editarlo.",
    "Completar los datos del formulario (nombre, dorsal, posicion, fecha de nacimiento, etc.).",
    "Hacer clic en 'Guardar'. Si el dorsal ya existe en el equipo, se muestra un aviso.",
]:
    bul(doc, paso)

body(doc,"Registro de partido:", bold=True)
for paso in [
    "Acceder al modulo 'Partidos' y hacer clic en 'Nuevo partido'.",
    "Completar los datos del partido (rival, fecha, tipo, competicion, formacion).",
    "En la pestana 'Estadisticas', introducir los datos colectivos del equipo.",
    "En la pestana 'Convocatoria', seleccionar los jugadores convocados y su rol (titular/suplente).",
    "En la pestana 'Eventos', registrar goles, tarjetas y cambios con el minuto exacto.",
    "Hacer clic en 'Guardar partido'.",
]:
    bul(doc, paso)

body(doc,"Generacion de informes PDF:", bold=True)
for paso in [
    "Acceder al modulo 'Informes'.",
    "Seleccionar el tipo de informe (plantilla, partido o temporada).",
    "Si es necesario, seleccionar el partido o el rango de fechas.",
    "Hacer clic en 'Generar PDF'. Se abre un selector para elegir la ubicacion de guardado.",
    "El PDF se genera automaticamente con los datos actualizados.",
]:
    bul(doc, paso)

h2(doc,"4.4. Manual de administracion","s4_4")
body(doc,"El perfil de directiva dispone de funcionalidades adicionales de administracion del sistema.")

body(doc,"Configuracion del club:", bold=True)
body(doc,"Desde el modulo 'Configuracion' se pueden modificar los datos del club (nombre, localidad, presidente, email de contacto, escudo y temporada activa). Cambiar la temporada activa afecta a los filtros de toda la aplicacion.")

body(doc,"Gestion de usuarios:", bold=True)
for paso in [
    "Acceder al modulo de 'Usuarios' desde el menu de administracion.",
    "Crear cuentas de usuario para entrenadores y jugadores indicando su correo y rol.",
    "Asignar entrenadores a equipos desde el modulo 'Equipos' -> 'Gestionar entrenadores'.",
    "Vincular jugadores a sus cuentas de usuario desde la ficha del jugador -> 'Vincular usuario'.",
    "Desactivar cuentas de usuario que ya no esten activas en el club.",
]:
    bul(doc, paso)

body(doc,"Gestion de equipos:", bold=True)
body(doc,"La directiva puede crear nuevos equipos para cada categoria y temporada, asignarles un campo habitual y su escudo. La eliminacion de un equipo es una operacion irreversible que debe confirmarse en el dialogo de aviso.")

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# 5. CONCLUSIONES
# ════════════════════════════════════════════════════════════════════════
h1(doc,"5. CONCLUSIONES FINALES","s5")
body(doc,"GRADO DE CUMPLIMIENTO DE LOS OBJETIVOS", bold=True)
make_table(doc,
    ["Objetivo","Estado","Observaciones"],
    [
        ["Aplicacion de escritorio JavaFX multiplataforma","Cumplido","Cliente instalable en Windows con interfaz completa en JavaFX 25"],
        ["API REST con Spring Boot en la nube","Cumplido","Desplegada en Google Cloud Run con 15+ endpoints documentados"],
        ["Base de datos PostgreSQL remota (Supabase)","Cumplido","13 tablas con relaciones implementadas y datos de prueba cargados"],
        ["Estadisticas y graficos de rendimiento","Cumplido","Dashboard con TilesFX y estadisticas acumuladas por jugador"],
        ["Generacion de informes PDF con JasperReports","Cumplido","3 tipos de informe: plantilla, partido y temporada"],
        ["Autenticacion JWT y cifrado BCrypt","Cumplido","Spring Security + JWT; BCrypt para contrasenas en el backend"],
        ["Sistema de roles (directiva/entrenador/jugador)","Cumplido","Control de acceso implementado en backend y en la UI del cliente"],
        ["Instalador nativo con jpackage","Cumplido","ScouteoSetup.exe generado e incluido en el repositorio"],
        ["Aplicacion movil complementaria","Parcialmente cumplido","Scouteo Mobile en desarrollo; funcionalidades basicas implementadas"],
    ],
    col_widths=[7,2.5,6])

body(doc,"PROPUESTA DE MODIFICACIONES Y AMPLIACIONES FUTURAS", bold=True)
for r in [
    "Modo offline completo: implementar una cache local (SQLite o Room) en el cliente que sincronice con Supabase cuando se recupere la conexion.",
    "Notificaciones push: integrar un sistema de notificaciones para avisar a los jugadores de convocatorias y cambios de entrenamiento.",
    "Modulo de scouting: permitir evaluar jugadores de otros equipos con fichas de observacion y comparativas de rendimiento.",
    "Integracion con wearables: importar datos de dispositivos de seguimiento fisico (frecuencia cardiaca, distancia recorrida) directamente a los entrenamientos.",
    "Panel web para directivos: desarrollar un dashboard web ligero que no requiera instalacion, para consulta rapida desde cualquier dispositivo.",
    "Inteligencia artificial: incorporar modelos de prediccion de rendimiento y sugerencia de alineaciones basados en el historico de partidos.",
    "Gestion economica del club: anadir un modulo de control de presupuestos, pagos de fichas y patrocinadores, acercando el sistema a un ERP deportivo completo.",
]:
    bul(doc, r)

body(doc,
    "En terminos globales, el proyecto ha cubierto los objetivos principales establecidos "
    "en el analisis inicial. La arquitectura cliente-servidor adoptada ha demostrado ser "
    "solida y escalable, y la separacion en capas facilita el mantenimiento y la extension "
    "futura del sistema. La implementacion ha requerido la aplicacion de conocimientos de "
    "todos los modulos del ciclo DAM, lo que demuestra la naturaleza verdaderamente "
    "intermodular del proyecto.")

page_break(doc)

# ════════════════════════════════════════════════════════════════════════
# 6. BIBLIOGRAFIA
# ════════════════════════════════════════════════════════════════════════
h1(doc,"6. BIBLIOGRAFIA","s6")
body(doc,"A continuacion se lista la bibliografia y webgrafia consultada durante el desarrollo del proyecto:")

for ref in [
    "[1] Oracle Corporation. (2024). Java SE 24 Documentation. https://docs.oracle.com/en/java/javase/24/",
    "[2] OpenJFX. (2024). JavaFX 25 API Documentation. https://openjfx.io/javadoc/25/",
    "[3] Pivotal Software / VMware. (2024). Spring Boot Reference Documentation 3.3. https://docs.spring.io/spring-boot/docs/3.3.x/reference/html/",
    "[4] Pivotal Software / VMware. (2024). Spring Data JPA Reference Documentation. https://docs.spring.io/spring-data/jpa/docs/current/reference/html/",
    "[5] Pivotal Software / VMware. (2024). Spring Security Reference. https://docs.spring.io/spring-security/reference/",
    "[6] Supabase, Inc. (2024). Supabase Documentation. https://supabase.com/docs",
    "[7] PostgreSQL Global Development Group. (2024). PostgreSQL 15 Documentation. https://www.postgresql.org/docs/15/",
    "[8] TIBCO Software / Cloud Software Group. (2024). JasperReports Library 7.0 Documentation. https://community.jaspersoft.com/documentation/",
    "[9] Philipp Wagner. (2023). favre-lib BCrypt 0.10.2. https://github.com/patrickfav/bcrypt",
    "[10] Hansolo (Gerrit Grunwald). (2024). TilesFX 21.0.9. https://github.com/HanSolo/tilesfx",
    "[11] ControlsFX Project. (2024). ControlsFX 11.2.1 API. https://controlsfx.github.io/",
    "[12] Google LLC. (2024). Gson 2.10.1. https://github.com/google/gson",
    "[13] Google Cloud. (2024). Cloud Run Documentation. https://cloud.google.com/run/docs",
    "[14] Oracle Corporation. (2024). jpackage - Package a self-contained Java application. https://docs.oracle.com/en/java/javase/24/docs/specs/man/jpackage.html",
    "[15] Gradle, Inc. (2024). Gradle User Manual 8.12. https://docs.gradle.org/8.12/userguide/userguide.html",
    "[16] Apache Maven Project. (2024). Maven Documentation. https://maven.apache.org/guides/",
    "[17] Auth0 by Okta. (2024). JSON Web Tokens. https://jwt.io/introduction",
    "[18] OpenAPI Initiative. (2024). OpenAPI Specification 3.1. https://spec.openapis.org/oas/v3.1.0",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    set_line_spacing(p)
    run = p.add_run(ref)
    run.font.name=FONT; run.font.size=Pt(11); run.font.color.rgb=C_NEGRO

# ─── GUARDAR ────────────────────────────────────────────────────────────
out = r"C:\Users\User\Desktop\Scouteo_inicios\DOCUMENTACION_SCOUTEO.docx"
doc.save(out)
print(f"OK -> {out}")
